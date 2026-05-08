from fastapi import APIRouter, Depends, HTTPException, Response
from sqlalchemy.orm import Session
from ..database import get_db
from ..models import Note
from ..auth import verify_clerk_token, ClerkUser
from fpdf import FPDF
from docx import Document
from io import BytesIO
import re

router = APIRouter(prefix="/api/notes", tags=["Notes Export"])

def clean_markdown(text):
    # Very basic markdown to plain text conversion
    text = re.sub(r'#+\s*(.*)', r'\1', text) # Headings
    text = re.sub(r'\*\*(.*)\*\*', r'\1', text) # Bold
    text = re.sub(r'\*(.*)\*', r'\1', text) # Italic
    text = re.sub(r'\[(.*)\]\(.*\)', r'\1', text) # Links
    text = re.sub(r'`(.*)`', r'\1', text) # Code
    return text

import urllib.parse

def sanitize_filename(name):
    # Remove non-ascii or special chars for header safety
    return urllib.parse.quote(name.replace(" ", "_"))

@router.get("/{note_id}/export/pdf")
async def export_note_pdf(
    note_id: str,
    db: Session = Depends(get_db),
    user: ClerkUser = Depends(verify_clerk_token)
):
    note = db.query(Note).filter(Note.id == note_id, Note.user_id == user.id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        
        # Safe encoding for FPDF1
        title_safe = note.title.encode('latin-1', 'replace').decode('latin-1')
        pdf.cell(0, 10, title_safe, ln=True)
        pdf.ln(5)
        
        pdf.set_font("Arial", "", 11)
        content = clean_markdown(note.content)
        content_safe = content.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 8, content_safe)
        
        pdf_bytes = pdf.output(dest='S')
        if isinstance(pdf_bytes, str):
            pdf_bytes = pdf_bytes.encode('latin-1')
            
        filename = sanitize_filename(note.title)
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

@router.get("/{note_id}/export/docx")
async def export_note_docx(
    note_id: str,
    db: Session = Depends(get_db),
    user: ClerkUser = Depends(verify_clerk_token)
):
    note = db.query(Note).filter(Note.id == note_id, Note.user_id == user.id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    try:
        doc = Document()
        doc.add_heading(note.title, 0)
        
        lines = note.content.split('\n')
        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.strip():
                doc.add_paragraph(clean_markdown(line))
                
        f = BytesIO()
        doc.save(f)
        f.seek(0)
        
        filename = sanitize_filename(note.title)
        return Response(
            content=f.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")

@router.get("/{note_id}/export/txt")
async def export_note_txt(
    note_id: str,
    db: Session = Depends(get_db),
    user: ClerkUser = Depends(verify_clerk_token)
):
    note = db.query(Note).filter(Note.id == note_id, Note.user_id == user.id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    content = f"Title: {note.title}\n\n{clean_markdown(note.content)}"
    filename = sanitize_filename(note.title)
    
    return Response(
        content=content,
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename={filename}.txt"}
    )
